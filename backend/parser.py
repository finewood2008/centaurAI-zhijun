"""文档解析器 — 支持 PDF/DOCX/PPTX/Excel/MD/TXT/图片/视频/音频"""
import hashlib
import logging
import re
import zipfile
import xml.etree.ElementTree as ET
from datetime import date, datetime, time
from pathlib import Path
from config import (
    SUPPORTED_TEXT_EXTENSIONS,
    SUPPORTED_IMAGE_EXTENSIONS,
    SUPPORTED_VIDEO_EXTENSIONS,
    SUPPORTED_AUDIO_EXTENSIONS,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
)

logger = logging.getLogger(__name__)


class EmptyFileError(ValueError):
    """The source exists but contains no data to parse or index."""


def parse_file(file_path: str) -> dict:
    """解析文件，返回 {text, metadata, file_type}；DOCX/PDF 额外返回可选 parts。

    parts 是结构化、可展示、可定位的派生部分：
      {
        "part_type": "paragraph" | "table" | "page",
        "ordinal": 1,
        "text": "可检索的纯文本（表格为保留行列边界的 TSV）",
        "location": {"page": 1} | {"section": 1, "table": 1}
      }
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    stat = path.stat()
    if stat.st_size == 0:
        raise EmptyFileError(f"文件为空，无法索引: {file_path}")

    metadata = {
        "file_path": str(path.absolute()),
        "file_name": path.name,
        "file_size": stat.st_size,
        "modified_time": stat.st_mtime,
    }

    if ext in SUPPORTED_TEXT_EXTENSIONS:
        text, parts = _parse_text_document_parts(path, ext)
        result = {"text": text, "metadata": metadata, "file_type": "text"}
        if parts:
            result["parts"] = parts
        return result

    elif ext in SUPPORTED_IMAGE_EXTENSIONS:
        return {"text": "", "metadata": metadata, "file_type": "image"}

    elif ext in SUPPORTED_VIDEO_EXTENSIONS:
        # 视频与图片一样保持轻量：此处只定类型，抽帧/转写/OCR 在 watcher.index_file 做
        return {"text": "", "metadata": metadata, "file_type": "video"}

    elif ext in SUPPORTED_AUDIO_EXTENSIONS:
        # 音频保持轻量：实际转写在 watcher.index_file 做
        return {"text": "", "metadata": metadata, "file_type": "audio"}

    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def _parse_text_document(path: Path, ext: str) -> str:
    return _parse_text_document_parts(path, ext)[0]


def _parse_text_document_parts(path: Path, ext: str) -> tuple[str, list[dict]]:
    """返回 (平铺文本, 结构化 parts)；仅 DOCX/PDF 产出 parts，其余返回空列表。"""
    if ext == ".pdf":
        return _parse_pdf_with_parts(path)
    elif ext == ".docx":
        return _parse_docx_with_parts(path)
    elif ext == ".pptx":
        return _parse_pptx(path), []
    elif ext in (".xlsx", ".xlsm", ".xls"):
        return _parse_excel(path, ext), []
    elif ext in (".md", ".txt"):
        return _parse_plain(path), []
    return "", []


def _normalize_cell(value) -> str:
    """把单元格内容压成单行可检索文本：None（空单元格）转空串，内部换行/制表符统一替换为空格。"""
    if value is None:
        return ""
    return str(value).replace("\r\n", " ").replace("\r", " ").replace("\n", " ").replace("\t", " ").strip()


def _image_dimensions(blob: bytes) -> tuple[int | None, int | None]:
    """从图片字节读取宽高；解析失败返回 (None, None)，不抛异常。"""
    try:
        from io import BytesIO
        from PIL import Image
        with Image.open(BytesIO(blob)) as image:
            return image.width, image.height
    except Exception:
        return None, None


def _parse_pdf(path: Path) -> str:
    return _parse_pdf_with_parts(path)[0]


def _parse_pdf_with_parts(path: Path) -> tuple[str, list[dict]]:
    """每页产出一个 page part；可识别表格时额外产出 table part。

    若当前 PyMuPDF 版本不支持表格识别，仅记录可诊断日志并返回空表格列表，
    不能让整份 PDF 解析失败。
    """
    import fitz
    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        raise ValueError(f"PDF 文件损坏: {path}: {exc}") from exc
    texts: list[str] = []
    parts: list[dict] = []
    ordinal = 0
    try:
        for page_index, page in enumerate(doc, start=1):
            page_text = page.get_text().strip()
            if page_text:
                texts.append(page_text)
                ordinal += 1
                parts.append({
                    "part_type": "page",
                    "ordinal": ordinal,
                    "text": page_text,
                    "location": {"page": page_index},
                })
            table_parts = _extract_pdf_tables(page, page_index)
            for table_part in table_parts:
                ordinal += 1
                table_part["ordinal"] = ordinal
                parts.append(table_part)
            image_parts = _extract_pdf_images(doc, page, page_index)
            for image_part in image_parts:
                ordinal += 1
                image_part["part_type"] = "image"
                image_part["ordinal"] = ordinal
                image_part["text"] = ""
                parts.append(image_part)
    finally:
        doc.close()
    return "\n\n".join(texts), parts


def _extract_pdf_images(doc, page, page_index: int) -> list[dict]:
    """从单页提取内嵌图片，每次实际放置各产出一个 image part。

    同一图片对象（xref）在同一页被放置多次时，用 page.get_image_rects(xref) 枚举每次
    放置位置，各自产出 part（location 带 occurrence / bbox）；无法定位放置位置时保留
    一次出现。此处不合并引用——文件级去重由持久化层按内容 hash 完成（同一 artifact
    文件、多个来源行，保留全部来源位置）。
    """
    images: list[dict] = []
    seen_xrefs: set[int] = set()
    try:
        page_images = page.get_images(full=True) or []
    except Exception as exc:
        logger.warning("PDF 图片清单读取失败（第 %d 页）: %s", page_index, exc)
        return images
    for item in page_images:
        xref = item[0]
        if xref in seen_xrefs:
            continue
        seen_xrefs.add(xref)
        try:
            info = doc.extract_image(xref)
        except Exception as exc:
            logger.warning("PDF 图片提取失败（第 %d 页 xref=%s）: %s", page_index, xref, exc)
            continue
        blob = info.get("image") if info else None
        if not blob:
            continue
        ext = str(info.get("ext") or "png").lower()
        try:
            rects = page.get_image_rects(xref) or []
        except Exception as exc:
            logger.warning("PDF 图片放置位置读取失败（第 %d 页 xref=%s）: %s", page_index, xref, exc)
            rects = []
        occurrences = rects or [None]
        for occurrence, rect in enumerate(occurrences, start=1):
            location: dict = {"page": page_index, "occurrence": occurrence}
            if rect is not None:
                location["bbox"] = [
                    float(rect.x0), float(rect.y0), float(rect.x1), float(rect.y1)
                ]
            images.append({
                "image": {
                    "blob": blob,
                    "mime": f"image/{ext}",
                    "width": info.get("width"),
                    "height": info.get("height"),
                },
                "location": location,
            })
    return images


def _extract_pdf_tables(page, page_index: int) -> list[dict]:
    """提取单页表格为 table parts（TSV 保留行列边界，空单元格保持空值）。"""
    table_parts: list[dict] = []
    try:
        finder = page.find_tables()
        tables = getattr(finder, "tables", []) or []
    except Exception as exc:
        logger.warning("PDF 表格识别暂不可用（第 %d 页）: %s", page_index, exc)
        return table_parts
    for table_index, table in enumerate(tables, start=1):
        try:
            rows = [
                [_normalize_cell(cell) for cell in row]
                for row in table.extract()
            ]
        except Exception as exc:
            logger.warning("PDF 表格提取失败（第 %d 页）: %s", page_index, exc)
            continue
        rows = [row for row in rows if any(cell for cell in row)]
        tsv = "\n".join("\t".join(row) for row in rows)
        if not tsv:
            continue
        table_parts.append({
            "part_type": "table",
            "ordinal": 0,  # 由调用方分配全局顺序号
            "text": tsv,
            "location": {"page": page_index, "table": table_index},
        })
    return table_parts


def _parse_docx(path: Path) -> str:
    return _parse_docx_with_parts(path)[0]


def _paragraph_text_full(paragraph) -> str:
    """段落全文：直接运行文本 + 段落内文本框(w:txbxContent)的嵌套文本。

    python-docx 的 Paragraph.text 只取 w:p 的直接子级 w:r 运行文本；文本框
    （文字方框）内的文字位于 w:pict / w:drawing 之下的 w:txbxContent，不会被
    提取，导致表单型 DOCX（如排期管理表）在详情页大量内容缺失。此函数补充
    提取文本框内全部段落，跨段以换行连接，并保持「直接文本 → 文本框文本」的
    阅读顺序。
    """
    from docx.oxml.ns import qn
    chunks: list[str] = []
    for txbx in paragraph._p.iter(qn("w:txbxContent")):
        lines: list[str] = []
        for p in txbx.iter(qn("w:p")):
            line = "".join(t.text or "" for t in p.iter(qn("w:t"))).strip()
            if line:
                lines.append(line)
        if lines:
            chunks.append("\n".join(lines))
    text = paragraph.text.strip()
    if not chunks:
        return text
    return "\n".join([text, *chunks]) if text else "\n".join(chunks)


def _cell_full_text(cell) -> str:
    """单元格全文：直接段落文本 + 单元格内文本框嵌套文本（不进入嵌套表格）。"""
    lines: list[str] = []
    for p in cell.paragraphs:
        text = _paragraph_text_full(p)
        if text:
            lines.append(text)
    return "\n".join(lines)


def _parse_docx_with_parts(path: Path) -> tuple[str, list[dict]]:
    """DOCX 段落 + 表格：按正文 body 的 w:p / w:tbl 原始顺序产出 part。

    段落为 paragraph part，表格为 TSV 的 table part。若先遍历全部段落再遍历全部
    表格，会丢失“段落 A → 表格 → 段落 B”的交错顺序，违反“按文档顺序排列”契约。
    """
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn
    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ValueError(f"DOCX 文件损坏: {path}: {exc}") from exc
    paragraphs: list[str] = []
    parts: list[dict] = []
    ordinal = 0
    table_index = 0
    para_ordinal = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, doc)
            text = _paragraph_text_full(paragraph)
            images = _extract_docx_paragraph_images(child, doc)
            if not text and not images:
                continue
            para_ordinal += 1
            if text:
                paragraphs.append(text)
                ordinal += 1
                parts.append({
                    "part_type": "paragraph",
                    "ordinal": ordinal,
                    "text": text,
                    "location": {},
                })
            for occurrence, image in enumerate(images, start=1):
                ordinal += 1
                image["part_type"] = "image"
                image["ordinal"] = ordinal
                image["text"] = ""
                image["location"] = {"paragraph": para_ordinal, "occurrence": occurrence}
                parts.append(image)
        elif child.tag == qn("w:tbl"):
            table_index += 1
            table = Table(child, doc)
            rows = []
            for row in table.rows:
                cells = [_normalize_cell(_cell_full_text(cell)) for cell in row.cells]
                if not any(cells):
                    continue
                rows.append(cells)
            tsv = "\n".join("\t".join(row) for row in rows)
            if tsv:
                ordinal += 1
                parts.append({
                    "part_type": "table",
                    "ordinal": ordinal,
                    "text": tsv,
                    "location": {"table": table_index},
                })
            # P14-02：表格单元格内的内嵌图片也要完整提取（复用同一关系解析逻辑）。
            for image in _extract_table_images(table, doc, table_index):
                ordinal += 1
                image["part_type"] = "image"
                image["ordinal"] = ordinal
                image["text"] = ""
                parts.append(image)
    return "\n".join(paragraphs), parts


def _extract_docx_paragraph_images(p_element, doc) -> list[dict]:
    """从单个 w:p 提取内嵌图片（inline drawing 的 a:blip + 旧版 VML imagedata）。

    同一 rId 被多次引用（同一段落内放置两次）时保留每次出现，各产出一个 part；
    文件级去重由持久化层按内容 hash 完成，此处不合并引用。
    """
    from docx.oxml.ns import qn
    # python-docx 的 nsmap 不含 VML 前缀，故用字面命名空间 URI。
    vml_imagedata = "{urn:schemas-microsoft-com:vml}imagedata"
    images: list[dict] = []
    for blip in p_element.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if not rid:
            continue
        image = _docx_image_from_rid(rid, doc)
        if image:
            images.append(image)
    for imagedata in p_element.iter(vml_imagedata):
        rid = imagedata.get(qn("r:id"))
        if not rid:
            continue
        image = _docx_image_from_rid(rid, doc)
        if image:
            images.append(image)
    return images


def _docx_image_from_rid(rid: str, doc) -> dict | None:
    """按关系 ID 取内嵌图片字节与元数据；读取失败返回 None 并记录日志。"""
    try:
        related = doc.part.related_parts[rid]
        blob = related.blob
    except (KeyError, AttributeError) as exc:
        logger.warning("DOCX 内嵌图片读取失败（rid=%s）: %s", rid, exc)
        return None
    if not blob:
        return None
    mime = str(getattr(related, "content_type", "") or "application/octet-stream")
    width, height = _image_dimensions(blob)
    return {"image": {"blob": blob, "mime": mime, "width": width, "height": height}}


def _extract_table_images(table, doc, table_index: int) -> list[dict]:
    """递归遍历表格单元格段落，提取其中的内嵌图片。

    位置记录 table / row / column（合并单元格按首次出现列计），同一单元格内多次出现
    以 occurrence 区分；文件级去重由持久化层按内容 hash 完成。
    """
    images: list[dict] = []
    for row_index, row in enumerate(table.rows, start=1):
        seen_cells: set[int] = set()
        for col_index, cell in enumerate(row.cells, start=1):
            tc_id = id(cell._tc)
            if tc_id in seen_cells:
                continue  # 合并单元格：只提取一次，位置取首次出现的列
            seen_cells.add(tc_id)
            cell_images: list[dict] = []
            for paragraph in cell.paragraphs:
                cell_images.extend(_extract_docx_paragraph_images(paragraph._p, doc))
            for occurrence, image in enumerate(cell_images, start=1):
                images.append({
                    **image,
                    "location": {
                        "table": table_index,
                        "row": row_index,
                        "column": col_index,
                        "occurrence": occurrence,
                    },
                })
    return images


def _parse_pptx(path: Path) -> str:
    """Extract visible slide text and speaker notes from a PPTX package."""
    texts: list[str] = []

    def sort_key(name: str) -> tuple[int, int]:
        match = re.search(r"(\d+)\.xml$", name)
        num = int(match.group(1)) if match else 0
        kind = 1 if "/notesSlides/" in name else 0
        return num, kind

    try:
        with zipfile.ZipFile(path) as zf:
            names = [
                n for n in zf.namelist()
                if re.match(r"ppt/slides/slide\d+\.xml$", n)
                or re.match(r"ppt/notesSlides/notesSlide\d+\.xml$", n)
            ]
            for name in sorted(names, key=sort_key):
                content = zf.read(name)
                page_text = _extract_pptx_xml_text(content)
                if page_text:
                    label = "备注" if "/notesSlides/" in name else "幻灯片"
                    num = re.search(r"(\d+)\.xml$", name)
                    heading = f"{label} {num.group(1)}" if num else label
                    texts.append(f"{heading}\n{page_text}")
    except zipfile.BadZipFile as e:
        raise ValueError(f"PPTX 文件损坏: {path}") from e

    return "\n\n".join(texts)


def _extract_pptx_xml_text(content: bytes) -> str:
    root = ET.fromstring(content)
    paragraphs: list[str] = []
    for para in root.iter("{http://schemas.openxmlformats.org/drawingml/2006/main}p"):
        runs: list[str] = []
        for text_node in para.iter("{http://schemas.openxmlformats.org/drawingml/2006/main}t"):
            if text_node.text:
                runs.append(text_node.text)
        text = "".join(runs).strip()
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs)


def _excel_cell_text(value) -> str:
    """Convert spreadsheet cell values to stable, searchable text."""
    if value is None:
        return ""
    if isinstance(value, bool):
        return "TRUE" if value else "FALSE"
    if isinstance(value, datetime):
        if value.time() == time.min:
            return value.date().isoformat()
        return value.isoformat(sep=" ", timespec="seconds")
    if isinstance(value, (date, time)):
        return value.isoformat()
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).replace("\r\n", " / ").replace("\r", " / ").replace("\n", " / ").replace("\t", " ").strip()


def _excel_sheet_text(title: str, rows) -> str:
    lines: list[str] = []
    for row in rows:
        cells = [_excel_cell_text(value) for value in row]
        while cells and not cells[-1]:
            cells.pop()
        if any(cells):
            lines.append("\t".join(cells))
    if not lines:
        return ""
    return f"工作表：{title}\n" + "\n".join(lines)


def _parse_excel(path: Path, ext: str) -> str:
    """Extract worksheet names and cell rows from modern and legacy Excel files."""
    try:
        if ext in (".xlsx", ".xlsm"):
            from openpyxl import load_workbook

            workbook = load_workbook(str(path), read_only=True, data_only=False)
            try:
                sheets = [
                    text
                    for sheet in workbook.worksheets
                    if (text := _excel_sheet_text(sheet.title, sheet.iter_rows(values_only=True)))
                ]
            finally:
                workbook.close()
            return "\n\n".join(sheets)

        import xlrd

        workbook = xlrd.open_workbook(str(path), on_demand=True)
        sheets: list[str] = []
        try:
            for sheet in workbook.sheets():
                rows = []
                for row_index in range(sheet.nrows):
                    values = []
                    for cell in sheet.row(row_index):
                        value = cell.value
                        if cell.ctype == xlrd.XL_CELL_DATE:
                            value = xlrd.xldate_as_datetime(value, workbook.datemode)
                        elif cell.ctype == xlrd.XL_CELL_BOOLEAN:
                            value = bool(value)
                        elif cell.ctype == xlrd.XL_CELL_ERROR:
                            value = xlrd.error_text_from_code.get(value, f"#ERROR({value})")
                        values.append(value)
                    rows.append(values)
                text = _excel_sheet_text(sheet.name, rows)
                if text:
                    sheets.append(text)
        finally:
            workbook.release_resources()
        return "\n\n".join(sheets)
    except Exception as e:
        raise ValueError(f"Excel 文件解析失败: {path.name}: {e}") from e


def _parse_plain(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def is_supported(file_path: str) -> bool:
    ext = Path(file_path).suffix.lower()
    return ext in (
        SUPPORTED_TEXT_EXTENSIONS
        | SUPPORTED_IMAGE_EXTENSIONS
        | SUPPORTED_VIDEO_EXTENSIONS
        | SUPPORTED_AUDIO_EXTENSIONS
    )


def file_hash(file_path: str) -> str:
    """文件内容 SHA1（增量索引用：内容未变则跳过重嵌）"""
    try:
        h = hashlib.sha1()
        with open(file_path, "rb") as f:
            for block in iter(lambda: f.read(65536), b""):
                h.update(block)
        return h.hexdigest()
    except OSError:
        return ""


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """将长文本切成带重叠的块，尽量在段落/句子边界切分。

    bge-small-zh 上限 512 token，整篇直接编码会被模型悄悄截断、丢失后文，
    分块后每块独立成向量，长文召回从「只看开头」变为「全文可命中」。
    """
    text = text.strip()
    if not text:
        return []
    if len(text) <= size:
        return [text]

    # 先按段落切，再贪心打包到 size 以内；超长段落按句子/硬切兜底
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    units: list[str] = []
    for para in paragraphs:
        if len(para) <= size:
            units.append(para)
        else:
            units.extend(_split_long(para, size))

    chunks: list[str] = []
    buf = ""
    for unit in units:
        if buf and len(buf) + len(unit) + 1 > size:
            chunks.append(buf)
            # 重叠：保留上一块尾部 overlap 个字符，维持跨块语义连续
            buf = (buf[-overlap:] + "\n" + unit) if overlap else unit
        else:
            buf = (buf + "\n" + unit) if buf else unit
    if buf:
        chunks.append(buf)
    return chunks


def _split_long(para: str, size: int) -> list[str]:
    """按中英文句子边界切超长段落，再不行就硬切"""
    sentences = re.split(r"(?<=[。！？.!?；;\n])", para)
    out: list[str] = []
    buf = ""
    for s in sentences:
        if not s:
            continue
        if len(s) > size:
            if buf:
                out.append(buf)
                buf = ""
            for i in range(0, len(s), size):
                out.append(s[i : i + size])
        elif len(buf) + len(s) > size:
            out.append(buf)
            buf = s
        else:
            buf += s
    if buf:
        out.append(buf)
    return out
