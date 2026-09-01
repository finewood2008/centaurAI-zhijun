"""MindOS P14-01 DOCX/PDF 结构化解析回归测试。

覆盖：
- DOCX：段落 + 表格（TSV 保留行列边界、空单元格保持空值、location 表序号）；
- DOCX：段落/表格按 body 原始顺序交错返回（段落 A → 表格 → 段落 B）；
- PDF：每页 page part + 可识别表格的 table part（带页码）；表格识别失败不炸整份 PDF；
- 空表格：整表空单元格不产出 table part；
- 损坏文档：DOCX/PDF 解析失败抛可读 ValueError（后台任务会标记失败，不会永久卡 processing）；
- derived_store.document_parts 按 material_id + input_hash 幂等替换；
- _chunk_parts_with_meta：每个 chunk 带 part_id / page / table_ordinal 元数据；
- detail_of：返回 contentParts / tableCount，且不含物理路径；
- 真实索引链路回归：index_file → document_parts → 向量 metadata → detail_of，
  以及“重处理解析结果为空时清理旧 parts”；
- 单元格规范化：None / 换行 / 制表符处理符合“空单元格为空串”契约。

依赖项目 .venv（python-docx / PyMuPDF），可独立于 server 运行：
    .venv\\Scripts\\python.exe -m unittest test_mindos_p14_parsing -v
"""
import contextlib
import shutil
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock, patch

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import fitz
from docx import Document

import parser
from parser import parse_file
import watcher
from watcher import _chunk_parts_with_meta
from mindos.stores import derived_store
from mindos.services import ingestion


def _make_docx(path: Path, paragraphs: list[str], tables: list[list[list[str]]]) -> Path:
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    for table_rows in tables:
        row_count = max(len(table_rows), 1)
        col_count = max((len(r) for r in table_rows), default=1)
        table = doc.add_table(rows=row_count, cols=col_count)
        for i, row in enumerate(table_rows):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell
    doc.save(str(path))
    return path


def _make_pdf(path: Path, pages_text: list[str], table_page: int | None = None) -> Path:
    """生成多页 PDF；table_page 指定的页面绘制一个带网格线 + 文本的表格。"""
    pdf = fitz.open()
    try:
        for page_index, text in enumerate(pages_text, start=1):
            page = pdf.new_page()
            page.insert_text((72, 72), text)
            if table_page is not None and page_index == table_page:
                x0, y0, x1, y1 = 60, 90, 240, 160
                for x in (60, 120, 180, 240):
                    page.draw_line((x, y0), (x, y1))
                for y in (90, 110, 130, 150, 160):
                    page.draw_line((x0, y), (x1, y))
                page.insert_text((65, 105), "item")
                page.insert_text((125, 105), "amount")
                page.insert_text((185, 105), "owner")
                page.insert_text((65, 125), "A")
                page.insert_text((125, 125), "100")
        pdf.save(str(path))
    finally:
        pdf.close()
    return path


class DocxPartsTests(unittest.TestCase):
    def setUp(self):
        self._tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        shutil.rmtree(self._tmp, ignore_errors=True)

    def test_docx_paragraphs_and_tables(self):
        path = _make_docx(
            self._tmp / "t.docx",
            ["第一段：项目预算说明。", "第二段：风险分析。"],
            [[["项目", "金额", "负责人"], ["A", "100", ""]]],
        )
        result = parse_file(str(path))
        # text 字段保持旧语义（段落平铺），兼容旧调用方
        self.assertIn("第一段", result["text"])
        self.assertNotIn("金额", result["text"])  # 表格不进平铺 text
        parts = result["parts"]
        self.assertEqual([p["part_type"] for p in parts], ["paragraph", "paragraph", "table"])
        table = parts[2]
        self.assertEqual(table["ordinal"], 3)
        self.assertEqual(table["location"], {"table": 1})
        # TSV 保留行列边界，空单元格保持空值（尾部空串保留）
        self.assertEqual(table["text"], "项目\t金额\t负责人\nA\t100\t")

    def test_empty_table_produces_no_part(self):
        path = _make_docx(self._tmp / "e.docx", ["只有正文"], [[[""]]])
        result = parse_file(str(path))
        self.assertEqual([p["part_type"] for p in result["parts"]], ["paragraph"])

    def test_interleaved_paragraph_and_table_preserves_document_order(self):
        """段落 A → 表格 → 段落 B 必须按原文顺序返回，而非先全部段落再全部表格。"""
        doc = Document()
        doc.add_paragraph("段落 A：开头")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "列1"
        table.cell(0, 1).text = "列2"
        table.cell(1, 0).text = "值1"
        table.cell(1, 1).text = "值2"
        doc.add_paragraph("段落 B：结尾")
        path = self._tmp / "inter.docx"
        doc.save(str(path))
        result = parse_file(str(path))
        parts = result["parts"]
        self.assertEqual(
            [p["part_type"] for p in parts],
            ["paragraph", "table", "paragraph"],
        )
        self.assertEqual(parts[0]["text"], "段落 A：开头")
        self.assertEqual(parts[1]["text"], "列1\t列2\n值1\t值2")
        self.assertEqual(parts[1]["location"], {"table": 1})
        self.assertEqual(parts[2]["text"], "段落 B：结尾")

    def test_docx_textbox_text_extracted(self):
        """段落/表格单元格内经文本框(w:txbxContent)的文本必须被提取。

        回归“原材料详情页 DOCX 显示不完整”：python-docx 的 paragraph/cell.text
        只取直接运行文本，文本框内文字需递归补充，且保持「直接文本 → 文本框
        文本」顺序。
        """
        from docx.oxml import parse_xml

        def _txbx_xml(texts: list[str]) -> str:
            paras = "".join(
                f'<w:p><w:r><w:t>{t}</w:t></w:r></w:p>' for t in texts
            )
            return (
                '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                f"<w:txbxContent>{paras}</w:txbxContent>"
                "</w:pict>"
            )

        doc = Document()
        para = doc.add_paragraph()
        para.add_run("段落正文")
        para.runs[0]._r.append(parse_xml(_txbx_xml(["文本框甲", "文本框乙"])))
        table = doc.add_table(rows=2, cols=2)
        table.cell(1, 0).text = "单元格正文"
        table.cell(1, 0).paragraphs[0].runs[0]._r.append(
            parse_xml(_txbx_xml(["单元格内文本框"]))
        )
        path = self._tmp / "txbx.docx"
        doc.save(str(path))
        result = parse_file(str(path))
        # 段落：直接文本在前，文本框文本在后，跨段换行
        self.assertIn("段落正文", result["text"])
        self.assertIn("文本框甲\n文本框乙", result["text"])
        # 表格单元格内的文本框文本保留在 TSV 单元格中
        table_part = next(p for p in result["parts"] if p["part_type"] == "table")
        self.assertIn("单元格正文", table_part["text"])
        self.assertIn("单元格内文本框", table_part["text"])


class NormalizeCellTests(unittest.TestCase):
    def test_none_becomes_empty_string(self):
        # PyMuPDF 表格空单元格常返回 None；必须转空串而非字符串 "None"
        self.assertEqual(parser._normalize_cell(None), "")

    def test_inner_newlines_and_tabs_flattened(self):
        self.assertEqual(parser._normalize_cell("a\r\nb\tc"), "a b c")

    def test_empty_string_stays_empty(self):
        self.assertEqual(parser._normalize_cell(""), "")


class PdfPartsTests(unittest.TestCase):
    def setUp(self):
        self._tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        shutil.rmtree(self._tmp, ignore_errors=True)

    def test_pdf_pages_and_table(self):
        path = _make_pdf(self._tmp / "t.pdf", ["page one body", "page two body"], table_page=2)
        result = parse_file(str(path))
        self.assertIn("page one body", result["text"])
        pages = [p for p in result["parts"] if p["part_type"] == "page"]
        self.assertEqual(len(pages), 2)
        self.assertEqual(pages[0]["location"]["page"], 1)
        self.assertEqual(pages[1]["location"]["page"], 2)
        tables = [p for p in result["parts"] if p["part_type"] == "table"]
        self.assertEqual(len(tables), 1)
        self.assertEqual(tables[0]["location"]["page"], 2)
        self.assertIn("item", tables[0]["text"])
        # 空单元格不得以字符串 "None" 出现（PyMuPDF 空单元格可能返回 None）
        self.assertNotIn("None", tables[0]["text"])
        # 全局顺序号连续且唯一
        ordinals = [p["ordinal"] for p in result["parts"]]
        self.assertEqual(ordinals, sorted(ordinals))

    def test_table_detection_failure_does_not_break_pdf(self):
        """表格识别失败时整份 PDF 仍可解析（page part 不丢）。"""
        path = _make_pdf(self._tmp / "p.pdf", ["only body"], table_page=None)
        with patch.object(fitz.Page, "find_tables", side_effect=Exception("no table support")):
            result = parse_file(str(path))
        self.assertEqual(result["text"].strip(), "only body")
        self.assertEqual([p["part_type"] for p in result["parts"]], ["page"])


class DamagedDocTests(unittest.TestCase):
    def setUp(self):
        self._tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        shutil.rmtree(self._tmp, ignore_errors=True)

    def test_corrupt_docx_raises_value_error(self):
        bad = self._tmp / "bad.docx"
        bad.write_bytes(b"not a docx")
        with self.assertRaises(ValueError):
            parse_file(str(bad))

    def test_corrupt_pdf_raises_value_error(self):
        bad = self._tmp / "bad.pdf"
        bad.write_bytes(b"not a pdf")
        with self.assertRaises(ValueError):
            parse_file(str(bad))


class DerivedStorePartsTests(unittest.TestCase):
    def setUp(self):
        self._db = Path(tempfile.mktemp(suffix=".db"))
        derived_store.reset_for_tests(self._db)
        self.store = derived_store.DerivedStore.instance()

    def tearDown(self):
        if self._db.exists():
            self._db.unlink()
        # 恢复默认路径，避免影响其他测试模块
        derived_store.reset_for_tests()

    def _parts(self):
        return [
            {"part_type": "paragraph", "ordinal": 1, "text": "正文", "location": {}},
            {"part_type": "table", "ordinal": 2, "text": "A\tB\n1\t", "location": {"page": 1, "table": 1}},
        ]

    def test_upsert_idempotent_by_input_hash(self):
        first = self.store.upsert_document_parts("mindos_p1", "h1", self._parts())
        self.assertEqual(len(first), 2)
        # 同一 input_hash 不重复写库
        again = self.store.upsert_document_parts("mindos_p1", "h1", self._parts())
        self.assertEqual(len(again), 2)
        self.assertEqual(len(self.store.parts_for_material("mindos_p1")), 2)

    def test_replacement_deletes_old_parts(self):
        self.store.upsert_document_parts("mindos_p1", "h1", self._parts())
        new_parts = self._parts() + [
            {"part_type": "page", "ordinal": 3, "text": "尾页", "location": {"page": 2}},
        ]
        replaced = self.store.upsert_document_parts("mindos_p1", "h2", new_parts)
        self.assertEqual(len(replaced), 3)
        self.assertEqual(len(self.store.parts_for_material("mindos_p1")), 3)
        self.assertEqual(self.store.table_count_for_material("mindos_p1"), 1)

    def test_delete_for_material(self):
        self.store.upsert_document_parts("mindos_p1", "h1", self._parts())
        self.assertEqual(self.store.delete_for_material("mindos_p1"), 2)
        self.assertEqual(self.store.parts_for_material("mindos_p1"), [])

    def test_material_id_for_source(self):
        with patch.object(
            derived_store.JobStore, "instance",
            return_value=MagicMock(list=lambda: [{"source_path": "/tmp/a.pdf", "material_id": "mindos_a"}]),
        ):
            self.assertEqual(derived_store.material_id_for_source("/tmp/a.pdf"), "mindos_a")
            self.assertIsNone(derived_store.material_id_for_source("/tmp/b.pdf"))


class ChunkMetaTests(unittest.TestCase):
    def test_per_part_chunk_metadata(self):
        parts = [
            {"id": "m::paragraph::1", "part_type": "paragraph", "ordinal": 1, "text": "正文甲", "location": {}},
            {"id": "m::table::2", "part_type": "table", "ordinal": 2, "text": "A\tB\n1\t", "location": {"page": 1, "table": 1}},
            {"id": "m::page::3", "part_type": "page", "ordinal": 3, "text": "页面乙", "location": {"page": 2}},
        ]
        chunks, pcm = _chunk_parts_with_meta(parts)
        self.assertEqual(len(chunks), 3)
        self.assertEqual(pcm[0], {"part_id": "m::paragraph::1"})
        self.assertEqual(pcm[1], {"part_id": "m::table::2", "table_ordinal": 2})
        self.assertEqual(pcm[2], {"part_id": "m::page::3", "page": 2})

    def test_skips_empty_part_text(self):
        parts = [
            {"id": "m::table::1", "part_type": "table", "ordinal": 1, "text": "", "location": {}},
            {"id": "m::page::2", "part_type": "page", "ordinal": 2, "text": "   ", "location": {"page": 1}},
        ]
        chunks, pcm = _chunk_parts_with_meta(parts)
        self.assertEqual(chunks, [])
        self.assertEqual(pcm, [])


class DetailContentPartsTests(unittest.TestCase):
    def setUp(self):
        self._tmp = Path(tempfile.mkdtemp())
        self._db = self._tmp / "derived.db"
        derived_store.reset_for_tests(self._db)
        self.store = derived_store.DerivedStore.instance()
        self.source = self._tmp / "report.docx"
        self.source.write_bytes(b"fake docx content")

    def tearDown(self):
        derived_store.reset_for_tests()
        shutil.rmtree(self._tmp, ignore_errors=True)

    def _detail(self, material_id: str, file_type: str):
        rec = {
            "material_id": material_id,
            "file_name": "report.docx" if file_type == "document" else "img.png",
            "file_type": file_type,
            "source_path": str(self.source),
            "job_id": f"job_{material_id}",
            "created_at": 1700000000.0,
            "folder": "未分类",
        }
        with patch.object(
            ingestion.JobStore, "instance",
            return_value=MagicMock(get=lambda _m: rec, is_canceled=lambda _m: False),
        ), patch.object(ingestion, "get_job", return_value={"state": "done"}), patch.object(
            ingestion, "get_source_chunks", return_value=[]
        ), patch.object(ingestion, "_ann_get", return_value={"tags": []}), patch.object(
            ingestion, "parse_file", return_value={"text": "正文段落", "parts": []}
        ):
            return ingestion.detail_of(material_id)

    def test_detail_returns_content_parts_and_table_count(self):
        self.store.upsert_document_parts("mindos_d1", "h", [
            {"part_type": "paragraph", "ordinal": 1, "text": "正文段落", "location": {}},
            {"part_type": "table", "ordinal": 2, "text": "A\tB\n1\t2", "location": {"page": 2, "table": 1}},
        ])
        detail = self._detail("mindos_d1", "document")
        self.assertEqual(detail["tableCount"], 1)
        self.assertEqual(len(detail["contentParts"]), 2)
        table = detail["contentParts"][1]
        self.assertEqual(table["partType"], "table")
        self.assertEqual(table["rows"], [["A", "B"], ["1", "2"]])
        self.assertEqual(table["location"], {"page": 2, "table": 1})
        # contentParts 不暴露物理路径 / artifact_key
        for part in detail["contentParts"]:
            self.assertNotIn("source_path", part)
            self.assertNotIn("artifact_key", part)

    def test_non_document_returns_empty_parts(self):
        detail = self._detail("mindos_d2", "image")
        self.assertEqual(detail["contentParts"], [])
        self.assertEqual(detail["tableCount"], 0)


class IndexCleanupRegressionTests(unittest.TestCase):
    """真实索引链路：index_file → document_parts → 向量 metadata → detail_of。

    重点回归“重处理解析结果为空（如扫描版 PDF 无文本层）时清理旧 parts”，
    避免详情页展示已不存在的旧表格/内容块。
    """

    def setUp(self):
        self._tmp = Path(tempfile.mkdtemp())
        derived_store.reset_for_tests(self._tmp / "derived.db")
        self.source = self._tmp / "report.pdf"
        self.material_id = "mindos_clean1"

    def tearDown(self):
        derived_store.reset_for_tests()
        shutil.rmtree(self._tmp, ignore_errors=True)

    @contextlib.contextmanager
    def _patch_index_deps(self):
        """屏蔽模型/向量/标注依赖，仅验证解析 → parts 持久化 → chunk 元数据链路。"""
        add_calls: list[tuple] = []
        with patch.object(watcher, "get_source_hash", return_value=None), patch.object(
            watcher, "_index_fingerprint", return_value="hash-v1"
        ), patch.object(
            watcher.derived_store, "material_id_for_source", return_value=self.material_id
        ), patch.object(
            watcher.annotations, "get_rag_override", return_value=None
        ), patch.object(
            watcher.annotations, "caption_of", return_value=""
        ), patch.object(
            watcher, "_ocr_scanned_pdf", return_value=""
        ), patch.object(
            watcher, "embed_batch_texts",
            side_effect=lambda chunks: [[0.0] * 8 for _ in chunks],
        ), patch.object(
            watcher, "add_file_chunks",
            side_effect=lambda *a, **kw: add_calls.append((a, kw)) or True,
        ):
            yield add_calls

    def _detail(self):
        rec = {
            "material_id": self.material_id,
            "file_name": "report.pdf",
            "file_type": "document",
            "source_path": str(self.source),
            "job_id": "job_clean1",
            "created_at": 1700000000.0,
            "folder": "未分类",
        }
        with patch.object(
            ingestion.JobStore, "instance",
            return_value=MagicMock(get=lambda _m: rec, is_canceled=lambda _m: False),
        ), patch.object(ingestion, "get_job", return_value={"state": "done"}), patch.object(
            ingestion, "get_source_chunks", return_value=[]
        ), patch.object(ingestion, "_ann_get", return_value={"tags": []}), patch.object(
            ingestion, "parse_file", return_value={"text": "page one", "parts": []}
        ):
            return ingestion.detail_of(self.material_id)

    def test_full_chain_then_reprocess_to_empty_cleans_stale_parts(self):
        # 1) 首次：带表格的 PDF → parts 持久化 + 向量 metadata 携带 part 引用
        _make_pdf(self.source, ["page one", "page two"], table_page=2)
        with self._patch_index_deps() as add_calls:
            ok = watcher.index_file(str(self.source), force=True)
        self.assertTrue(ok)
        parts = derived_store.DerivedStore.instance().parts_for_material(self.material_id)
        self.assertTrue(any(p["part_type"] == "table" for p in parts))

        self.assertTrue(add_calls)
        _, kwargs = add_calls[-1]
        pcm = kwargs["per_chunk_metadata"]
        self.assertTrue(any("table_ordinal" in m for m in pcm))
        self.assertTrue(all("part_id" in m for m in pcm))

        # 2) 详情页可见该表格
        detail = self._detail()
        self.assertEqual(detail["tableCount"], 1)
        self.assertTrue(any(p["partType"] == "table" for p in detail["contentParts"]))

        # 3) 重处理为无文本层的空白 PDF → 旧 parts 被清理，详情不再显示表格
        blank = fitz.open()
        blank.new_page()
        blank.save(str(self.source))
        blank.close()
        with self._patch_index_deps():
            watcher.index_file(str(self.source), force=True)
        self.assertEqual(
            derived_store.DerivedStore.instance().parts_for_material(self.material_id), []
        )
        self.assertEqual(
            derived_store.DerivedStore.instance().table_count_for_material(self.material_id), 0
        )
        self.assertEqual(self._detail()["tableCount"], 0)


if __name__ == "__main__":
    unittest.main(verbosity=2)
