"""MindOS P14-02 内嵌图片提取、OCR 与预览回归测试。

覆盖：
- DOCX：从 word/media 提取内嵌图片（blob / mime / 尺寸 / 来源段落）；
- PDF：按页提取内嵌图片（来源页码）；
- 去重：同一图片多处引用 → 落盘单文件、多 part 行共享同一 artifact_key；
- derived_store：图片落盘、替换后孤儿文件清理、删除时清理图片目录；
- 受控读取：image_file_path 越界（路径穿越）拒绝；
- 索引链路：index_file → 图片落盘 → OCR 回写（ok/empty/unavailable）→
  chunk 元数据 modality=embedded_image_ocr + part_id → detail_of.embeddedImages；
- 受控读取接口：mindos_material_part_file 只允许本资料的 image part。

依赖项目 .venv（python-docx / PyMuPDF / PIL），可独立于 server 运行：
    .venv\\Scripts\\python.exe -m unittest test_mindos_p14_images -v
"""
import contextlib
import shutil
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock, patch

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import fitz
from docx import Document
from fastapi.responses import FileResponse

import parser
import watcher
from mindos.stores import derived_store
from mindos.services import ingestion
from mindos import uploads


def _make_png_bytes(size=(40, 30), color=(20, 60, 140)) -> bytes:
    from io import BytesIO
    from PIL import Image
    buf = BytesIO()
    Image.new("RGB", size, color).save(buf, format="PNG")
    return buf.getvalue()


def _make_png_file(path: Path, size=(40, 30), color=(20, 60, 140)) -> Path:
    from PIL import Image
    Image.new("RGB", size, color).save(str(path), format="PNG")
    return path


class DocxImageExtractionTests(unittest.TestCase):
    def setUp(self):
        self._tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        shutil.rmtree(self._tmp, ignore_errors=True)

    def test_docx_embedded_image_part(self):
        doc = Document()
        doc.add_paragraph("图片前的段落")
        png = _make_png_file(self._tmp / "a.png")
        doc.add_picture(str(png))
        doc.add_paragraph("图片后的段落")
        path = self._tmp / "img.docx"
        doc.save(str(path))

        result = parser.parse_file(str(path))
        images = [p for p in result["parts"] if p["part_type"] == "image"]
        self.assertEqual(len(images), 1)
        image = images[0]
        self.assertEqual(image["location"], {"paragraph": 2, "occurrence": 1})
        self.assertEqual(image["image"]["mime"], "image/png")
        self.assertEqual(image["image"]["width"], 40)
        self.assertEqual(image["image"]["height"], 30)
        self.assertGreater(len(image["image"]["blob"]), 0)

    def test_docx_table_cell_image_part(self):
        """表格单元格内的内嵌图片必须被完整提取，并记录 table / row / column。"""
        png = _make_png_file(self._tmp / "a.png")
        doc = Document()
        doc.add_paragraph("正文段落")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "a"
        table.cell(0, 1).text = "b"
        cell = table.cell(1, 0).paragraphs[0]
        cell.add_run().add_picture(str(png))
        path = self._tmp / "table-img.docx"
        doc.save(str(path))

        result = parser.parse_file(str(path))
        images = [p for p in result["parts"] if p["part_type"] == "image"]
        self.assertEqual(len(images), 1)
        self.assertEqual(
            images[0]["location"],
            {"table": 1, "row": 2, "column": 1, "occurrence": 1},
        )
        self.assertEqual(images[0]["image"]["mime"], "image/png")

    def test_docx_same_paragraph_multiple_occurrences_kept(self):
        """同一段落内同图出现两次 → 保留两次引用（occurrence 1/2），文件去重交给持久化层。"""
        png = _make_png_file(self._tmp / "a.png")
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run().add_picture(str(png))
        paragraph.add_run().add_picture(str(png))
        path = self._tmp / "same-paragraph.docx"
        doc.save(str(path))

        result = parser.parse_file(str(path))
        images = [p for p in result["parts"] if p["part_type"] == "image"]
        self.assertEqual(len(images), 2)
        self.assertEqual(
            [img["location"] for img in images],
            [{"paragraph": 1, "occurrence": 1}, {"paragraph": 1, "occurrence": 2}],
        )
        # 两次出现内容一致，持久化层将共用同一 artifact 文件
        self.assertEqual(images[0]["image"]["blob"], images[1]["image"]["blob"])


class PdfImageExtractionTests(unittest.TestCase):
    def setUp(self):
        self._tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        shutil.rmtree(self._tmp, ignore_errors=True)

    def test_pdf_embedded_image_part(self):
        png = _make_png_file(self._tmp / "a.png")
        pdf = fitz.open()
        try:
            page1 = pdf.new_page()
            page1.insert_text((72, 72), "first page")
            page2 = pdf.new_page()
            page2.insert_text((72, 72), "second page")
            page2.insert_image(fitz.Rect(60, 100, 160, 130), filename=str(png))
            path = self._tmp / "img.pdf"
            pdf.save(str(path))
        finally:
            pdf.close()

        result = parser.parse_file(str(path))
        images = [p for p in result["parts"] if p["part_type"] == "image"]
        self.assertEqual(len(images), 1)
        self.assertEqual(images[0]["location"]["page"], 2)
        self.assertEqual(images[0]["location"]["occurrence"], 1)
        self.assertIn("bbox", images[0]["location"])
        self.assertEqual(images[0]["image"]["mime"], "image/png")

    def test_same_image_on_multiple_pages_keeps_each_location(self):
        png = _make_png_file(self._tmp / "a.png")
        pdf = fitz.open()
        try:
            for _ in range(2):
                page = pdf.new_page()
                page.insert_image(fitz.Rect(60, 100, 160, 130), filename=str(png))
            path = self._tmp / "multi.pdf"
            pdf.save(str(path))
        finally:
            pdf.close()

        result = parser.parse_file(str(path))
        images = [p for p in result["parts"] if p["part_type"] == "image"]
        self.assertEqual(len(images), 2)
        self.assertEqual({p["location"]["page"] for p in images}, {1, 2})
        # 两页图片内容一致（blob 相同），文件级去重由持久化层完成
        self.assertEqual(
            images[0]["image"]["blob"], images[1]["image"]["blob"]
        )

    def test_same_image_twice_on_same_page_keeps_each_occurrence(self):
        """同一页内同图放置两次 → 保留两次引用（occurrence 1/2），文件去重交给持久化层。"""
        png = _make_png_file(self._tmp / "a.png")
        pdf = fitz.open()
        try:
            page = pdf.new_page()
            page.insert_image(fitz.Rect(60, 100, 160, 130), filename=str(png))
            page.insert_image(fitz.Rect(200, 100, 300, 130), filename=str(png))
            path = self._tmp / "same-page.pdf"
            pdf.save(str(path))
        finally:
            pdf.close()

        result = parser.parse_file(str(path))
        images = [p for p in result["parts"] if p["part_type"] == "image"]
        self.assertEqual(len(images), 2)
        occurrences = sorted(img["location"]["occurrence"] for img in images)
        self.assertEqual(occurrences, [1, 2])
        # 两次出现内容一致，持久化层将共用同一 artifact 文件
        self.assertEqual(images[0]["image"]["blob"], images[1]["image"]["blob"])


class DerivedStoreImageTests(unittest.TestCase):
    def setUp(self):
        self._tmp = Path(tempfile.mkdtemp())
        derived_store.reset_for_tests(self._tmp / "derived.db")
        self._imgdir = self._tmp / "images"
        self._imgdir.mkdir()
        self._patch = patch.object(derived_store, "DERIVED_IMAGES_DIR", self._imgdir)
        self._patch.start()
        self.store = derived_store.DerivedStore.instance()

    def tearDown(self):
        self._patch.stop()
        derived_store.reset_for_tests()
        shutil.rmtree(self._tmp, ignore_errors=True)

    def _image_part(self, ordinal, blob, location):
        return {
            "part_type": "image",
            "ordinal": ordinal,
            "text": "",
            "location": location,
            "image": {"blob": blob, "mime": "image/png", "width": 40, "height": 30},
        }

    def test_same_blob_dedup_to_single_file(self):
        blob = _make_png_bytes()
        parts = [
            self._image_part(1, blob, {"page": 1}),
            self._image_part(2, blob, {"page": 2}),
        ]
        rows = self.store.upsert_document_parts("mindos_img1", "h1", parts)
        self.assertEqual(len(rows), 2)
        keys = {row["artifact_key"] for row in rows}
        self.assertEqual(len(keys), 1)
        self.assertEqual(len(list((self._imgdir / "mindos_img1").iterdir())), 1)
        self.assertEqual(rows[0]["image_meta"]["mime"], "image/png")
        self.assertEqual(rows[0]["image_meta"]["ocr_status"], "empty")

    def test_replace_cleans_orphan_image_file(self):
        blob_a = _make_png_bytes(color=(10, 20, 30))
        blob_b = _make_png_bytes(color=(40, 50, 60))
        self.store.upsert_document_parts(
            "mindos_img1", "h1", [self._image_part(1, blob_a, {"page": 1})]
        )
        self.assertEqual(len(list((self._imgdir / "mindos_img1").iterdir())), 1)
        self.store.upsert_document_parts(
            "mindos_img1", "h2", [self._image_part(1, blob_b, {"page": 1})]
        )
        files = list((self._imgdir / "mindos_img1").iterdir())
        self.assertEqual(len(files), 1)
        self.assertEqual(files[0].read_bytes(), blob_b)

    def test_delete_removes_image_dir(self):
        self.store.upsert_document_parts(
            "mindos_img1", "h1", [self._image_part(1, _make_png_bytes(), {"page": 1})]
        )
        self.assertTrue((self._imgdir / "mindos_img1").is_dir())
        self.store.delete_for_material("mindos_img1")
        self.assertFalse((self._imgdir / "mindos_img1").exists())

    def test_image_file_path_rejects_traversal(self):
        self.assertIsNone(self.store.image_file_path("mindos_img1", "../secret.png"))
        self.assertIsNone(self.store.image_file_path("mindos_img1", "a/b.png"))
        self.assertIsNone(self.store.image_file_path("mindos_img1", "..\\x.png"))
        rows = self.store.upsert_document_parts(
            "mindos_img1", "h1", [self._image_part(1, _make_png_bytes(), {"page": 1})]
        )
        path = self.store.image_file_path("mindos_img1", rows[0]["artifact_key"])
        self.assertIsNotNone(path)
        self.assertTrue(path.is_file())


class IndexOcrIntegrationTests(unittest.TestCase):
    """index_file → 图片落盘 → OCR 回写 → chunk 元数据 → detail_of.embeddedImages。"""

    def setUp(self):
        self._tmp = Path(tempfile.mkdtemp())
        derived_store.reset_for_tests(self._tmp / "derived.db")
        self._imgdir = self._tmp / "images"
        self._imgdir.mkdir()
        self._patch = patch.object(derived_store, "DERIVED_IMAGES_DIR", self._imgdir)
        self._patch.start()
        self.source = self._tmp / "report.pdf"
        self.material_id = "mindos_img2"

    def tearDown(self):
        self._patch.stop()
        derived_store.reset_for_tests()
        shutil.rmtree(self._tmp, ignore_errors=True)

    @contextlib.contextmanager
    def _patch_deps(self, ocr_text: str, ocr_ready: bool):
        add_calls: list[tuple] = []
        with patch.object(watcher, "get_source_hash", return_value=None), patch.object(
            watcher, "_index_fingerprint", return_value="hash-v1"
        ), patch.object(
            watcher.derived_store, "material_id_for_source", return_value=self.material_id
        ), patch.object(
            watcher.annotations, "get_rag_override", return_value=None
        ), patch.object(
            watcher.annotations, "caption_of", return_value=""
        ), patch.object(
            watcher, "_ocr_scanned_pdf", return_value=""
        ), patch.object(
            watcher, "OCR_ENABLED", True
        ), patch.object(
            watcher, "ocr_available", return_value=ocr_ready
        ), patch.object(
            watcher, "ocr_image", return_value=ocr_text
        ), patch.object(
            watcher, "embed_batch_texts",
            side_effect=lambda chunks: [[0.0] * 8 for _ in chunks],
        ), patch.object(
            watcher, "add_file_chunks",
            side_effect=lambda *a, **kw: add_calls.append((a, kw)) or True,
        ):
            yield add_calls

    def _make_pdf_with_image(self):
        png = _make_png_file(self._tmp / "a.png")
        pdf = fitz.open()
        try:
            page = pdf.new_page()
            page.insert_text((72, 72), "first page text")
            page.insert_image(fitz.Rect(60, 100, 160, 130), filename=str(png))
            pdf.save(str(self.source))
        finally:
            pdf.close()

    def _image_part(self):
        parts = derived_store.DerivedStore.instance().parts_for_material(self.material_id)
        images = [p for p in parts if p["part_type"] == "image"]
        self.assertEqual(len(images), 1)
        return images[0]

    def _detail(self):
        rec = {
            "material_id": self.material_id,
            "file_name": "report.pdf",
            "file_type": "document",
            "source_path": str(self.source),
            "job_id": "job_img2",
            "created_at": 1700000000.0,
            "folder": "未分类",
        }
        with patch.object(
            ingestion.JobStore, "instance",
            return_value=MagicMock(get=lambda _m: rec, is_canceled=lambda _m: False),
        ), patch.object(ingestion, "get_job", return_value={"state": "done"}), patch.object(
            ingestion, "get_source_chunks", return_value=[]
        ), patch.object(ingestion, "_ann_get", return_value={"tags": []}), patch.object(
            ingestion, "parse_file", return_value={"text": "first page text", "parts": []}
        ):
            return ingestion.detail_of(self.material_id)

    def test_index_ocr_ok_and_chunk_metadata(self):
        self._make_pdf_with_image()
        with self._patch_deps(ocr_text="预算 2026", ocr_ready=True) as add_calls:
            ok = watcher.index_file(str(self.source), force=True)
        self.assertTrue(ok)
        image = self._image_part()
        self.assertEqual(image["text"], "预算 2026")
        self.assertEqual(image["image_meta"]["ocr_status"], "ok")
        # 向量 metadata：OCR 文本块带 modality=embedded_image_ocr + part_id
        _, kwargs = add_calls[-1]
        pcm = kwargs["per_chunk_metadata"]
        ocr_meta = [m for m in pcm if m.get("modality") == "embedded_image_ocr"]
        self.assertTrue(ocr_meta)
        self.assertTrue(all("part_id" in m for m in ocr_meta))
        # detail_of 返回 embeddedImages（含受控 previewUrl / ocrText / ocrStatus）
        detail = self._detail()
        matches = [i for i in detail["embeddedImages"] if i["partId"] == image["id"]]
        self.assertEqual(len(matches), 1)
        self.assertEqual(matches[0]["ocrText"], "预算 2026")
        self.assertEqual(matches[0]["ocrStatus"], "ok")
        self.assertTrue(matches[0]["previewUrl"].endswith(f"/parts/{image['id']}/file"))
        self.assertEqual(matches[0]["location"]["page"], 1)

    def test_ocr_engine_unavailable_sets_status(self):
        self._make_pdf_with_image()
        with self._patch_deps(ocr_text="", ocr_ready=False):
            watcher.index_file(str(self.source), force=True)
        image = self._image_part()
        self.assertEqual(image["image_meta"]["ocr_status"], "unavailable")
        self.assertEqual(image["text"], "")

    def test_ocr_empty_sets_status(self):
        self._make_pdf_with_image()
        with self._patch_deps(ocr_text="", ocr_ready=True):
            watcher.index_file(str(self.source), force=True)
        image = self._image_part()
        self.assertEqual(image["image_meta"]["ocr_status"], "empty")
        self.assertEqual(image["text"], "")


class PartFileEndpointTests(unittest.TestCase):
    """受控读取接口：只允许本资料的 image part，越界 / 归属不符一律 404。"""

    def setUp(self):
        self._tmp = Path(tempfile.mkdtemp())
        derived_store.reset_for_tests(self._tmp / "derived.db")
        self._imgdir = self._tmp / "images"
        self._imgdir.mkdir()
        self._patch = patch.object(derived_store, "DERIVED_IMAGES_DIR", self._imgdir)
        self._patch.start()
        self.store = derived_store.DerivedStore.instance()
        self.source = self._tmp / "report.pdf"

    def tearDown(self):
        self._patch.stop()
        derived_store.reset_for_tests()
        shutil.rmtree(self._tmp, ignore_errors=True)

    def _job_store_mock(self):
        return patch.object(
            ingestion.JobStore, "instance",
            return_value=MagicMock(get=lambda m: {
                "material_id": m, "file_name": "report.pdf", "file_type": "document",
                "source_path": str(self.source),
            }),
        )

    def test_image_part_returns_fileresponse(self):
        rows = self.store.upsert_document_parts(
            "mindos_end1", "h1",
            [{"part_type": "image", "ordinal": 1, "text": "", "location": {"page": 1},
              "image": {"blob": _make_png_bytes(), "mime": "image/png", "width": 40, "height": 30}}],
        )
        with self._job_store_mock():
            resp = uploads.mindos_material_part_file("mindos_end1", rows[0]["id"])
        self.assertIsInstance(resp, FileResponse)

    def test_non_image_part_404(self):
        rows = self.store.upsert_document_parts(
            "mindos_end1", "h1",
            [{"part_type": "page", "ordinal": 1, "text": "x", "location": {"page": 1}}],
        )
        with self._job_store_mock():
            with self.assertRaises(Exception) as ctx:
                uploads.mindos_material_part_file("mindos_end1", rows[0]["id"])
        self.assertEqual(getattr(ctx.exception, "status_code", None), 404)

    def test_part_belonging_to_other_material_404(self):
        rows = self.store.upsert_document_parts(
            "mindos_end1", "h1",
            [{"part_type": "image", "ordinal": 1, "text": "", "location": {"page": 1},
              "image": {"blob": _make_png_bytes(), "mime": "image/png"}}],
        )
        # 用另一资料 ID 访问 → 归属校验失败 → 404
        with self._job_store_mock():
            with self.assertRaises(Exception) as ctx:
                uploads.mindos_material_part_file("mindos_end2", rows[0]["id"])
        self.assertEqual(getattr(ctx.exception, "status_code", None), 404)

    def test_missing_part_404(self):
        with self._job_store_mock():
            with self.assertRaises(Exception) as ctx:
                uploads.mindos_material_part_file("mindos_end1", "no-such-part")
        self.assertEqual(getattr(ctx.exception, "status_code", None), 404)


if __name__ == "__main__":
    unittest.main(verbosity=2)
